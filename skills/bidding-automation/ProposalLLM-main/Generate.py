import openpyxl
import docx
import openai
import os
import requests
import json
import re  # Importing regular expressions for cleaning textexitt
from docx.shared import Inches
import shutil
from docx import Document
import io
import os
from docx.shared import Cm

#
#这个具有读取ExcelB列，如果有内容则标题二+1，没有内容标题三+1，自动根据C列生成D列的应答，对应章节号自动写入F列，将对应标准需求写入到Word里，不改变文字和图片
#原始的word 1- XXX ，2- XXX
#如果标题名称不同，需要手工设置last_heading_2，last_heading_3，具体参见 def get_last_section_numbers(doc)，用于填写Excel内容


#--------------------------------需要调整的全局变量--------------------------------------------------



# 设置OpenAI API Key
openai.api_key = "自己的Open AIKey"
# 设置百度 API Key
API_KEY = "自己百度的API Key"
SECRET_KEY = "自己百度的密钥"


# 设置Word输出最大宽度为14厘米
MAX_WIDTH_CM = 14.0
#对于点对点应答的Prompt
Prompt_Answer="现在有个问答，比选文件要求和比选申请人应答，我给你举个例子 ，比如比选文件要求：支持可视化创建不同类型数据源，包括但不限于：传统数据库、文件系统、消息队列、SaaS API，NoSQL等、必选申请人回答的是：完全支持。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，支持搜索功能。你学习一下我的风格。现在我是比选申请人，请严格按照我的风格来回答，请注意我回答的格式：首先是'完全支持'，然后说'系统支持什么什么', 这个过程需要你按照问题回答，不要跑题。例如，输入我的整体回答就变成了：'完全支持。系统支持数据源配置化管理，数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置高级管理功能，全面支持搜索功能。'以下是输入文字：" 
#对于内容应答的Prompt
Prompt_Content="你是一个大数据平台的专业产品售前，请针对这一需求给出800字的产品功能介绍，不要开头和总结，直接写产品功能，不需要用markdown格式，直接文本格式+特殊项目符号输出即可，需求如下:" 

#针对产品同样的内容重写标书内容，以应对不同版本标书的Prompt
Prompt_RewriteContent="" 
#针对不同需求，缩写需求内容，变为每个章节的小标题
Prompt_Title="你是一个专业作者，请把以下这段文字变为10字以内不带细节内容和标点和解释的文字，直接给出结果不要'简化为'这种返回：" 


MoreSection = 1  # 当为1时，启用新的生成标题策略，可以直接生成3级标题并读取 def get_last_section_numbers(doc)
ReGenerateText = 0  # 当为1时，会对原始产品文档当中的文字内容进行重写
DDDAnswer = 1 # 当为1时，会生成点对点应答
key_flag = 0 # 当为1时，★ ▲ 会自动带到标题和需求描述当中
level1 = 'heading 1'  
level2 = 'heading 2'
last_heading_1 = 2 #技术标书起始的段落，Word当中应该有第一章，概述，那么技术从第二大章开始，此处就是2
last_heading_2 = 0
last_heading_3 = 0

# 设置输出文本内容为空
content_between_headings = []
#-----------------------------------------------------------------------------------------------------------------------

# 百度调用Token
def get_access_token():
    """
    使用 AK，SK 生成鉴权签名（Access Token）
    :return: access_token，或是None(如果错误)
    """
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    return str(requests.post(url, params=params).json().get("access_token"))



# 函数：调用百度生成1000字的解决方案
def generate_solution_from_chatgpt(c_content):
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()

    prompt=  f"{Prompt_Content} {c_content}"
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content":f"'{prompt}'"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    response = requests.request("POST", url, headers=headers, data=payload)
    data = json.loads(response.text.strip())
    result_value = data['result']
    print('----生成内容:',result_value)
    return result_value



# 函数：调用ChatGPT API 生成1000字的解决方案
def generate_solution_from_chatgpt2(c_content):
    try:
        prompt = f"{Prompt_Content} {c_content}"
        response = openai.Completion.create(
            engine="gpt-4",
            prompt=prompt,
            max_tokens=1500,  # 1000字的最大token限制
            temperature=0.7
        )
        solution = response.choices[0].text.strip()
        return solution
    except Exception as e:
        print(f"Error generating solution: {e}")
        return "Error generating solution."
    


# 函数：调用百度 API，重写Word中的文本内容
def rewrite_word_content(c_column_content, word_content):
     
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()

    prompt= f"{Prompt_RewriteContent}你是一个专业售前，在写一个标书，目前在写对应 {c_column_content} 的需求，请你重写 {word_content} 这段文字"
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content":f"'{prompt}'"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    response = requests.request("POST", url, headers=headers, data=payload)
    data = json.loads(response.text.strip())
    result_value = data['result']
    print('----内容:',cleaned_title)
    return result_value

 
# 函数：调用ChatGPT API重写Word中的文本内容
def rewrite_word_content2(c_column_content, word_content):
    try:
        prompt = f"{Prompt_RewriteContent}你是一个专业售前，在写一个标书，目前在写对应 {c_column_content} 的需求，请你重写 {word_content} 这段文字"
        response = openai.Completion.create(
            engine="gpt-4",
            prompt=prompt,
            max_tokens=500,
            temperature=0.7
        )
        rewritten_text = response.choices[0].text.strip()
        return rewritten_text
    except Exception as e:
        print(f"Error rewriting word content: {e}")
        return word_content  # 如果API调用失败，返回原始内容
    

# 函数：调用百度 API，将文本缩减为15字以内的标题
def shorten_text(text):
    
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content":f"{Prompt_Title}'{text}'"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    response = requests.request("POST", url, headers=headers, data=payload)
    data = json.loads(response.text.strip())
    result_value = data['result']
    cleaned_title = result_value.replace("。", "")  # 去除中文句号“。”
     # 如果KeyFlag为1，处理★和▲符号
    if key_flag == 1:
            if '★' in text and '★' not in cleaned_title:
                cleaned_title = f"★{cleaned_title}"
            elif '▲' in text and '▲' not in cleaned_title:
                cleaned_title = f"▲{cleaned_title}"
    #print(last_heading_2，'.','----需求:',text,'标题变为：',cleaned_title)
    return cleaned_title

    
# 函数：调用ChatGPT API，将文本缩减为15字以内的标题
def shorten_text2(text):
       
    try:
        response = openai.Completion.create(
            engine="gpt-4",  # 选择模型
            prompt=f"Simplify the following content into a title of 15 characters or less: {text}",
            max_tokens=15,
            temperature=0.5
        )
        return response.choices[0].text.strip()
    except Exception as e:
        print(f"Error shortening text: {e}")
        return text[:15]  # 出错时返回前15个字

# 函数：调用百度API，优化需求说明满足需求条件
def optimize_description(text):

        
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token=" + get_access_token()
    
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content":f"{Prompt_Answer}'{text}'"
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    
    response = requests.request("POST", url, headers=headers, data=payload)
    data = json.loads(response.text.strip())
    result_value = data['result']

    #print('应答需求:',text,'应答:',result_value)

    return result_value



# 函数：调用ChatGPT API，优化需求说明为100字左右
def optimize_description2(text):
    try:
        response = openai.Completion.create(
            engine="gpt-4",
            prompt=f"重新书写这段话，变为100字左右的说明: {text}",
            max_tokens=200,
            temperature=0.7
        )
        optimized_text = response.choices[0].text.strip()
        return optimized_text
    except Exception as e:
        print(f"Error optimizing text: {e}")
        return text  # 返回原始文本以防API调用失败

# 将G列内容转换为整数，作为要打开的Word文档名，支持X-某文字.docx格式
def find_word_file(g_column_value):
    try:
        # 从G列获取值并转换为整数
        if g_column_value is not None and g_column_value != "":
            x_value = g_column_value
        else:
             x_value = ""  # 如果G列为空，返回“”

        # 查找匹配的文件，忽略文件名中"-"之后的文字
        for file_name in os.listdir():  # 假设所有文件在当前目录下
            match = re.match(rf"^{x_value}-.*\.docx$", file_name)
            if match:
                return file_name

        # 如果没有找到匹配的文件
        print(f"No file found for pattern: {x_value}-*.docx")
        return "no-file"

    except ValueError:
        print(f"Error converting G column to integer: {g_column_value}")
        return "no-file"

##################################################################
# Save content as new Word document 
def save_content_to_new_doc(content,  target_doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for item in content:
        if isinstance(item, str):
            paragraph = target_doc.add_paragraph()
            run = paragraph.add_run(item)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif isinstance(item, tuple) and item[0] == 'table':
            table_data = item[1]
            table = target_doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            tbl = table._element
            tbl_pr = tbl.tblPr
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tbl_borders.append(border)
            tbl_pr.append(tbl_borders)
        elif isinstance(item, tuple) and item[0] == 'image':
            image_stream = item[1]
            image_name = item[2]
            width_cm = item[3]
            height_cm = item[4]
            if width_cm > MAX_WIDTH_CM:
                scale_factor = MAX_WIDTH_CM / width_cm
                width_cm = MAX_WIDTH_CM
                height_cm = height_cm * scale_factor
            target_doc.add_paragraph().add_run().add_picture(image_stream, width=Cm(width_cm), height=Cm(height_cm))

        elif isinstance(item, tuple) and item[0] == 'list':
            # Add list paragraph to the new document while retaining its list format
            paragraph = target_doc.add_paragraph(item[1], style='List Paragraph')
    # 每个章节添加两个换行
    target_doc.add_paragraph()  # 第一个换行
    target_doc.add_paragraph()  # 第二个换行
    # new_doc.save(file_name)


# Function to extract images from runs
def get_image_from_run(run):
    drawing_elements = run._element.xpath('.//a:blip')
    
    if drawing_elements:
        embed = drawing_elements[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image_part = run.part.related_parts[embed]
        image_data = image_part.blob
        drawing_element = run._element.xpath('.//wp:extent')
        if drawing_element:
            cx = int(drawing_element[0].get('cx'))
            cy = int(drawing_element[0].get('cy'))
            width_cm = cx / 914400 * 2.54
            height_cm = cy / 914400 * 2.54
        else:
            width_cm, height_cm = None, None

        image_stream = io.BytesIO(image_data)
        image_stream.name = os.path.basename(image_part.partname)
        return image_stream, image_stream.name, width_cm, height_cm
    
    return None

# Function to iterate over document elements in order
def iter_block_items(parent):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.document import Document

    parent_elm = parent.element.body if isinstance(parent, Document) else parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


##################################################################


# 函数：从Word文档中获取最新的标题二和标题三章节号
def get_last_section_numbers(doc):
    last_heading_2 = 0
    last_heading_3 = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 2':  # 查找标题二
            match = re.match(r"(\d+)", paragraph.text)
            if match:
                last_heading_2 = int(match.group(1))
                last_heading_3 = 0  # 新的标题二，重置标题三的章节号

        elif paragraph.style.name == 'Heading 3':  # 查找标题三
            match = re.match(r"(\d+)\.(\d+)", paragraph.text)
            if match:
                last_heading_2 = int(match.group(1))  # 获取当前标题二章节号
                last_heading_3 = int(match.group(2))  # 获取当前标题三章节号
    
    #如果标题名称不同，需要手工设置last_heading_2，last_heading_3
    return last_heading_2, last_heading_3


# 函数：复制文本和图片，并保留图片在原文中的位置
def copy_content_with_images(doc, target_doc):
    from docx import Document
    from docx.text.paragraph import Paragraph 
    from docx.table import Table
    from docx.oxml.ns import qn

    content_between_headings = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para = block
            if para.style.name.startswith('Heading'):
                if current_level is not None:
                    save_content_to_new_doc(content_between_headings,target_doc)
                    content_between_headings = []
                heading_level = para.style.name.split(' ')[-1]
                current_level = heading_level
                current_heading_text = para.text
                content_between_headings.append(para.text)
            else:
                paragraph_text = ""
                for run in para.runs:
                    image_data = get_image_from_run(run)
                    if image_data:
                        image_stream, image_name, width_cm, height_cm = image_data
                        content_between_headings.append(('image', image_stream, image_name, width_cm, height_cm))
                    else:
                        text = run.text.strip()
                        if text:
                            paragraph_text += text

                # Check if paragraph is a list (ordered or unordered) and copy it with its format
                if para.style.name in ['List Paragraph'] or para._element.xpath('.//w:numPr'):
                    # Set the paragraph text and ensure all runs use the '宋体' font
                    list_text = para.text
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置东亚字体为宋体
                    content_between_headings.append(('list', list_text))
                elif paragraph_text:
                    # Ensure non-list paragraphs also use '宋体'
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    content_between_headings.append(paragraph_text)
    
        elif isinstance(block, Table):
            table = block
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ''
                    for cell_paragraph in cell.paragraphs:
                        for cell_run in cell_paragraph.runs:
                            image_stream = get_image_from_run(cell_run)
                            if image_stream:
                                pass
                            else:
                                cell_text += cell_run.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            content_between_headings.append(('table', table_data))
    
    if content_between_headings:
        save_content_to_new_doc(content_between_headings,target_doc)
      


# 打开Excel文件
excel_file = "需求对应表.xlsx"
wb = openpyxl.load_workbook(excel_file)
sheet = wb.active

# 打开目标Word文档
word_file = "标书内容.docx"
doc = docx.Document(word_file)

# 获取Word文档中的最后标题二和标题三章节号
last_heading_2, last_heading_3 = get_last_section_numbers(doc)

count=2
# 从第二行开始遍历Excel内容
for row in sheet.iter_rows(min_row=2):
    b_column_content = row[1].value  # 获取B列内容
    c_column_content = row[2].value  # 获取C列内容
    g_column_value = row[6].value  # 获取G列内容
    print(count,'------>',c_column_content)
    count = count+1
    # 判断行是否为空，如果为空则停止
    if not c_column_content:
        break

    # 调用ChatGPT API，缩减C列内容
    shortened_title = shorten_text(c_column_content)
    # 调用ChatGPT API，对C列做应答
    optimized_description = optimize_description(c_column_content)



    # 将优化后的说明写入E列
    row[4].value = optimized_description
    #将需求写入到D列
    row[3].value = shortened_title

    # 根据MoreSection逻辑生成标题，并处理标题二和标题三的级联关系
    if MoreSection == 1:
        if b_column_content:  # 如果B列有内容，生成标题二
            last_heading_2 += 1  # 标题二递增
            last_heading_3 = 1  # 新的标题二，重置标题三
            doc.add_heading(f" {b_column_content}", level=2)
            
            # 将当前标题二章节号写入F列
            print(f"当前章节{last_heading_1}.{last_heading_2}.{last_heading_3}")
            row[5].value = f"{last_heading_1}.{last_heading_2}.{last_heading_3}"

            # 生成标题三的小节，标题基于C列内容
            
            doc.add_heading(f"{shortened_title}", level=3)
        else:  # 如果B列没有内容，生成标题三
            last_heading_3 += 1  # 标题三递增
            doc.add_heading(f" {shortened_title}", level=3)

            # 将当前标题三章节号写入F列
            print(f"当前章节{last_heading_1}.{last_heading_2}.{last_heading_3}")
            row[5].value = f"{last_heading_1}.{last_heading_2}.{last_heading_3}"
    else:
        # 按原始逻辑生成标题二
        last_heading_2 += 1  # 标题二递增
        doc.add_heading(f"{last_heading_2}. {shortened_title}", level=2)

        # 将当前标题二章节号写入F列
        row[5].value = f"{last_heading_1}.{last_heading_2}"
        print(f"当前章节{last_heading_1}.{last_heading_2}")


    # 将G列内容转换为整数，作为要打开的Word文档名
    try:
        x_word_file = find_word_file(g_column_value)
        # 根据DDDAnswer的逻辑生成点对点应答内容
        if DDDAnswer == 1:
            # 插入C列的内容作为第一段
              doc.add_paragraph(c_column_content, style='Normal')
        
             # 插入"答："加上D列生成的内容作为第二段
              doc.add_paragraph(f"答：{optimized_description}", style='Normal')
        print('file=',g_column_value)
        # 检查对应的Word文件是否存在
        if os.path.exists(x_word_file):
            x_doc = docx.Document(x_word_file)
            # copy文件和图片

            copy_content_with_images(x_doc, doc)
        else:
        # 如果没有找到文件，调用ChatGPT API 生成解决方案 
             solution_text = generate_solution_from_chatgpt(c_column_content)
             doc.add_paragraph(solution_text, style='Normal')
    

    except ValueError:
        print(f"Error converting G column to integer: {g_column_value}")
        continue



# 保存更新后的Excel文件
wb.save(excel_file)

# 保存最终Word文档
doc.save(word_file)

print("Process completed!")